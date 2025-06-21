import streamlit as st
st.set_page_config(
    page_title="Terra Price Estimator",
    page_icon="🧠",  # or replace with image: "favicon.png"
)
# Groq model pricing (input/output in $ per 1M tokens)
PRICING = {
    "Llama 4 Scout (17Bx16E)":       {"in": 0.11, "out": 0.34},
    "Llama 4 Maverick (17Bx128E)":   {"in": 0.20, "out": 0.60},
    "Llama Guard 4 12B 128k":        {"in": 0.20, "out": 0.20},
    "DeepSeek R1 Distill Llama 70B": {"in": 0.75, "out": 0.99},
    "Qwen3 32B 131k":                {"in": 0.29, "out": 0.59},
    "Qwen QwQ 32B (Preview)":        {"in": 0.29, "out": 0.39},
    "Mistral Saba 24B":              {"in": 0.79, "out": 0.79},
    "Llama 3.3 70B Versatile 128k":  {"in": 0.59, "out": 0.79},
    "Llama 3.1 8B Instant 128k":     {"in": 0.05, "out": 0.08},
    "Llama 3 70B 8k":                {"in": 0.59, "out": 0.79},
    "Llama 3 8B 8k":                 {"in": 0.05, "out": 0.08},
    "Gemma 2 9B 8k":                 {"in": 0.20, "out": 0.20},
    "Llama Guard 3 8B 8k":           {"in": 0.20, "out": 0.20},
}

# App title
st.title("🧠 Groq Token Cost Estimator")

# Model selection
# Create a mapping with formatted labels
formatted_model_labels = {
    model: f"{model} (${info['in']} in / ${info['out']} out)"
    for model, info in PRICING.items()
}

# Reverse map to get actual model key from the display name
reverse_lookup = {v: k for k, v in formatted_model_labels.items()}

# Show display names in multiselect
model_labels = st.multiselect(
    "Select one or more models (with pricing):",
    options=list(formatted_model_labels.values()),
    default=[
        formatted_model_labels["Llama 3.3 70B Versatile 128k"],
        formatted_model_labels["Qwen3 32B 131k"]
    ]
)

# Resolve selected labels back to actual model keys
models = [reverse_lookup[label] for label in model_labels]


# Q&A input
st.header("🔍 Q&A Usage")
tokens_per_question = st.number_input("Tokens per question", value=10000, min_value=1)
users = st.number_input("Number of users", value=3, min_value=1)
questions_per_user = st.number_input("Questions per user/day", value=20, min_value=1)

# Report input
st.header("📄 Report Generation")
questions_per_report = st.number_input("Questions per report", value=30, min_value=1)
reports_per_day = st.number_input("Reports per day", value=2, min_value=1)

# Total token usage
total_questions_daily = users * questions_per_user
total_qna_tokens = tokens_per_question * total_questions_daily
total_report_tokens = questions_per_report * tokens_per_question * reports_per_day
total_daily_tokens = total_qna_tokens + total_report_tokens
monthly_tokens = total_daily_tokens * 30

st.markdown(f"**Total tokens/day:** {total_daily_tokens:,}")
st.markdown(f"**Total tokens/month:** {monthly_tokens:,}")

# Token allocation sliders
st.subheader("⚖️ Token Allocation (%)")
model_weights = {}
total_weight = 0

for model in models:
    weight = st.slider(f"{model}", min_value=0, max_value=100, value=int(100/len(models)), step=1)
    model_weights[model] = weight
    total_weight += weight

# Style for output
st.markdown("""
<style>
.big-money {
    font-size: 1.4em;
    font-weight: 600;
}
.model-cost {
    margin-bottom: 10px;
    font-size: 1.2em;
}
.combined {
    font-size: 1.3em;
    font-weight: 700;
    padding-top: 10px;
}
</style>
""", unsafe_allow_html=True)
if total_weight != 100:
    st.error(f"Total allocation must be 100%. Current total: {total_weight}%")
else:
    st.markdown("## 💵 <span class='big-money'>Cost Estimates</span>", unsafe_allow_html=True)

    qna_total_cost = 0.0
    report_total_cost = 0.0

    st.markdown("### 💬 Normal Q&A Cost Breakdown")
    for model in models:
        price = PRICING[model]
        model_share = total_qna_tokens * (model_weights[model] / 100)
        input_tokens = model_share * 0.85
        output_tokens = model_share * 0.15

        cost_input = (input_tokens * price["in"]) / 1_000_000
        cost_output = (output_tokens * price["out"]) / 1_000_000
        daily_cost = cost_input + cost_output
        monthly_cost = daily_cost * 30
        qna_total_cost += daily_cost

        st.markdown(f"""
        <div class='model-cost'>
        • <b>{model}</b>: <b>${daily_cost:.2f}</b>/day → <b>${monthly_cost:.2f}</b>/month
        </div>
        """, unsafe_allow_html=True)

    st.markdown("### 📄 Report Generation Cost Breakdown")
    for model in models:
        price = PRICING[model]
        model_share = total_report_tokens * (model_weights[model] / 100)
        input_tokens = model_share * 0.85
        output_tokens = model_share * 0.15

        cost_input = (input_tokens * price["in"]) / 1_000_000
        cost_output = (output_tokens * price["out"]) / 1_000_000
        daily_cost = cost_input + cost_output
        monthly_cost = daily_cost * 30
        report_total_cost += daily_cost

        st.markdown(f"""
        <div class='model-cost'>
        • <b>{model}</b>: <b>${daily_cost:.2f}</b>/day → <b>${monthly_cost:.2f}</b>/month
        </div>
        """, unsafe_allow_html=True)

    combined_daily_cost = qna_total_cost + report_total_cost
    combined_monthly_cost = combined_daily_cost * 30

    st.markdown(f"""
    <div class='combined'>🏁 Final Combined Total</div>
    <ul>
        <li><b>Q&A Daily:</b> <span class='big-money'>${qna_total_cost:.2f}</span></li>
        <li><b>Report Daily:</b> <span class='big-money'>${report_total_cost:.2f}</span></li>
        <li><b>Total Daily:</b> <span class='big-money'>${combined_daily_cost:.2f}</span></li>
        <li><b>Total Monthly:</b> <span class='big-money'>${combined_monthly_cost:.2f}</span></li>
    </ul>
    """, unsafe_allow_html=True)

from docx import Document

from io import BytesIO

# Create Word document
doc = Document()
doc.add_heading("Terra Price Estimator Report", 0)

doc.add_heading("Selected Models and Cost Breakdown", level=1)

for model in models:
    weight = model_weights[model]
    price = PRICING[model]
    model_share = total_daily_tokens * (weight / 100)
    input_tokens = model_share * 0.85
    output_tokens = model_share * 0.15
    cost_input = (input_tokens * price["in"]) / 1_000_000
    cost_output = (output_tokens * price["out"]) / 1_000_000
    daily_cost = cost_input + cost_output
    monthly_cost = daily_cost * 30

        # Q&A cost
    qna_tokens_model = total_qna_tokens * (weight / 100)
    qna_input = qna_tokens_model * 0.85
    qna_output = qna_tokens_model * 0.15
    qna_cost = (qna_input * price["in"] + qna_output * price["out"]) / 1_000_000

    # Report cost
    report_tokens_model = total_report_tokens * (weight / 100)
    rep_input = report_tokens_model * 0.85
    rep_output = report_tokens_model * 0.15
    rep_cost = (rep_input * price["in"] + rep_output * price["out"]) / 1_000_000

    doc.add_paragraph(
        f"{model}\n"
        f"  - Q&A Daily: ${qna_cost:.2f}\n"
        f"  - Reports Daily: ${rep_cost:.2f}\n"
        f"  - Combined Daily: ${qna_cost + rep_cost:.2f}, Monthly: ${(qna_cost + rep_cost) * 30:.2f}"
    )

doc.add_heading("Summary", level=1)
doc.add_paragraph(f"Total Q&A Tokens per Day: {total_qna_tokens:,}")
doc.add_paragraph(f"Total Report Tokens per Day: {total_report_tokens:,}")
doc.add_paragraph(f"Combined Daily Tokens: {total_daily_tokens:,}")
doc.add_paragraph(f"Q&A Daily Cost: ${qna_total_cost:.2f}")
doc.add_paragraph(f"Report Daily Cost: ${report_total_cost:.2f}")
doc.add_paragraph(f"Total Daily Cost: ${combined_daily_cost:.2f}")
doc.add_paragraph(f"Total Monthly Cost: ${combined_monthly_cost:.2f}")

# Export to in-memory buffer
buffer = BytesIO()
doc.save(buffer)
buffer.seek(0)

# Download button
st.download_button(
    label="📄 Download Word Report",
    data=buffer,
    file_name="terra_price_estimate.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)